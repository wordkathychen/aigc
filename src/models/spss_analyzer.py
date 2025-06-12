from typing import Dict, List, Any, Optional
import pandas as pd
import numpy as np
from scipy import stats
from sklearn.decomposition import FactorAnalysis
from sklearn.cluster import KMeans
from utils.exceptions import AnalysisError
from utils.logger import setup_logger

logger = setup_logger(__name__)

class SPSSAnalyzer:
    def __init__(self, data: pd.DataFrame):
        self.data = data
        
    def descriptive_analysis(self, columns: List[str]) -> Dict[str, Any]:
        """描述性统计分析"""
        try:
            results = {}
            for col in columns:
                desc = self.data[col].describe()
                results[col] = {
                    'mean': desc['mean'],
                    'std': desc['std'],
                    'min': desc['min'],
                    'max': desc['max'],
                    'median': desc['50%'],
                    'skewness': stats.skew(self.data[col]),
                    'kurtosis': stats.kurtosis(self.data[col])
                }
            return results
        except Exception as e:
            logger.error(f"描述性统计分析失败: {str(e)}")
            raise AnalysisError(f"无法完成描述性统计分析: {str(e)}")
    
    def t_test(self, 
              variable: str, 
              group_var: str,
              alpha: float = 0.05) -> Dict[str, Any]:
        """独立样本t检验"""
        try:
            group1 = self.data[self.data[group_var] == 0][variable]
            group2 = self.data[self.data[group_var] == 1][variable]
            
            t_stat, p_value = stats.ttest_ind(group1, group2)
            
            return {
                't_statistic': t_stat,
                'p_value': p_value,
                'significant': p_value < alpha,
                'mean_diff': group1.mean() - group2.mean()
            }
        except Exception as e:
            logger.error(f"t检验分析失败: {str(e)}")
            raise AnalysisError(f"无法完成t检验分析: {str(e)}")
    
    def correlation_analysis(self, 
                          variables: List[str],
                          method: str = 'pearson') -> pd.DataFrame:
        """相关性分析"""
        try:
            corr_matrix = self.data[variables].corr(method=method)
            p_values = pd.DataFrame(
                [[stats.pearsonr(self.data[v1], self.data[v2])[1] 
                  for v2 in variables] for v1 in variables],
                columns=variables,
                index=variables
            )
            
            return {
                'correlation': corr_matrix,
                'p_values': p_values
            }
        except Exception as e:
            logger.error(f"相关性分析失败: {str(e)}")
            raise AnalysisError(f"无法完成相关性分析: {str(e)}")
    
    def regression_analysis(self,
                         dependent: str,
                         independents: List[str]) -> Dict[str, Any]:
        """回归分析"""
        try:
            X = self.data[independents]
            y = self.data[dependent]
            
            # 添加常数项
            X = sm.add_constant(X)
            
            # 拟合模型
            model = sm.OLS(y, X).fit()
            
            return {
                'summary': model.summary(),
                'r_squared': model.rsquared,
                'adj_r_squared': model.rsquared_adj,
                'coefficients': model.params,
                'p_values': model.pvalues,
                'conf_int': model.conf_int()
            }
        except Exception as e:
            logger.error(f"回归分析失败: {str(e)}")
            raise AnalysisError(f"无法完成回归分析: {str(e)}")
    
    def factor_analysis(self, variables: List[str], n_factors: int = 3) -> Dict[str, Any]:
        """因子分析"""
        try:
            # 提取数据
            X = self.data[variables]
            
            # 创建因子分析模型
            fa = FactorAnalysis(n_components=n_factors, random_state=42)
            fa.fit(X)
            
            # 获取因子载荷
            loadings = pd.DataFrame(
                fa.components_.T,
                columns=[f'Factor{i+1}' for i in range(n_factors)],
                index=variables
            )
            
            # 计算解释方差比
            variance_ratio = fa.explained_variance_ratio_
            
            return {
                'loadings': loadings,
                'variance_ratio': variance_ratio,
                'total_variance': sum(variance_ratio)
            }
        except Exception as e:
            logger.error(f"因子分析失败: {str(e)}")
            raise AnalysisError(f"无法完成因子分析: {str(e)}")
    
    def anova_analysis(self, 
                      dependent: str, 
                      group_var: str) -> Dict[str, Any]:
        """单因素方差分析"""
        try:
            # 获取所有组
            groups = [group for _, group in self.data.groupby(group_var)[dependent]]
            
            # 执行方差分析
            f_stat, p_value = stats.f_oneway(*groups)
            
            # 计算组间描述统计
            group_stats = self.data.groupby(group_var)[dependent].describe()
            
            # 执行事后检验 (Tukey's HSD)
            from statsmodels.stats.multicomp import pairwise_tukeyhsd
            tukey = pairwise_tukeyhsd(self.data[dependent], self.data[group_var])
            
            return {
                'f_statistic': f_stat,
                'p_value': p_value,
                'group_statistics': group_stats,
                'post_hoc': tukey
            }
        except Exception as e:
            logger.error(f"方差分析失败: {str(e)}")
            raise AnalysisError(f"无法完成方差分析: {str(e)}")
    
    def cluster_analysis(self, 
                        variables: List[str], 
                        n_clusters: int = 3) -> Dict[str, Any]:
        """聚类分析"""
        try:
            # 数据标准化
            from sklearn.preprocessing import StandardScaler
            scaler = StandardScaler()
            X = scaler.fit_transform(self.data[variables])
            
            # K-means聚类
            kmeans = KMeans(n_clusters=n_clusters, random_state=42)
            clusters = kmeans.fit_predict(X)
            
            # 计算每个簇的中心点
            centers = pd.DataFrame(
                scaler.inverse_transform(kmeans.cluster_centers_),
                columns=variables
            )
            
            # 计算轮廓系数
            from sklearn.metrics import silhouette_score
            silhouette_avg = silhouette_score(X, clusters)
            
            # 添加聚类标签到原数据
            results_df = self.data.copy()
            results_df['Cluster'] = clusters
            
            # 计算每个簇的样本数
            cluster_sizes = pd.Series(clusters).value_counts().sort_index()
            
            return {
                'cluster_labels': clusters,
                'cluster_centers': centers,
                'silhouette_score': silhouette_avg,
                'cluster_sizes': cluster_sizes,
                'results': results_df
            }
        except Exception as e:
            logger.error(f"聚类分析失败: {str(e)}")
            raise AnalysisError(f"无法完成聚类分析: {str(e)}")
    
    def chi_square_test(self,
                       var1: str,
                       var2: str,
                       alpha: float = 0.05) -> Dict[str, Any]:
        """卡方检验"""
        try:
            # 创建列联表
            contingency_table = pd.crosstab(self.data[var1], self.data[var2])
            
            # 执行卡方检验
            chi2, p_value, dof, expected = stats.chi2_contingency(contingency_table)
            
            # 计算克拉默系数 (Cramer's V)
            n = contingency_table.sum().sum()
            min_dim = min(contingency_table.shape) - 1
            cramer_v = np.sqrt(chi2 / (n * min_dim))
            
            return {
                'contingency_table': contingency_table,
                'chi2_statistic': chi2,
                'p_value': p_value,
                'degrees_of_freedom': dof,
                'expected_frequencies': pd.DataFrame(
                    expected,
                    index=contingency_table.index,
                    columns=contingency_table.columns
                ),
                'cramers_v': cramer_v,
                'significant': p_value < alpha
            }
        except Exception as e:
            logger.error(f"卡方检验失败: {str(e)}")
            raise AnalysisError(f"无法完成卡方检验: {str(e)}")

    def export_results(self, results: Dict[str, Any], output_path: str) -> None:
        """导出分析结果到Excel，支持所有分析类型"""
        try:
            with pd.ExcelWriter(output_path) as writer:
                # 导出描述性统计结果
                if 'descriptive' in results:
                    pd.DataFrame(results['descriptive']).to_excel(
                        writer, 
                        sheet_name='描述性统计'
                    )
                
                # 导出相关性分析结果
                if 'correlation' in results:
                    results['correlation']['correlation'].to_excel(
                        writer,
                        sheet_name='相关性分析_相关系数'
                    )
                    results['correlation']['p_values'].to_excel(
                        writer,
                        sheet_name='相关性分析_P值'
                    )
                
                # 导出回归分析结果
                if 'regression' in results:
                    summary_df = pd.DataFrame({
                        'R方': [results['regression']['r_squared']],
                        '调整R方': [results['regression']['adj_r_squared']]
                    })
                    summary_df.to_excel(writer, sheet_name='回归分析_汇总')
                    
                    coef_df = pd.DataFrame({
                        '系数': results['regression']['coefficients'],
                        'P值': results['regression']['p_values']
                    })
                    coef_df.to_excel(writer, sheet_name='回归分析_系数')
                
                # 导出因子分析结果
                if 'factor' in results:
                    results['factor']['loadings'].to_excel(
                        writer,
                        sheet_name='因子分析_载荷'
                    )
                    pd.Series(
                        results['factor']['variance_ratio'],
                        name='解释方差比'
                    ).to_excel(writer, sheet_name='因子分析_方差')
                
                # 导出方差分析结果
                if 'anova' in results:
                    pd.DataFrame({
                        'F统计量': [results['anova']['f_statistic']],
                        'P值': [results['anova']['p_value']]
                    }).to_excel(writer, sheet_name='方差分析_结果')
                    results['anova']['group_statistics'].to_excel(
                        writer,
                        sheet_name='方差分析_描述统计'
                    )
                
                # 导出卡方检验结果
                if 'chi_square' in results:
                    results['chi_square']['contingency_table'].to_excel(
                        writer,
                        sheet_name='卡方检验_列联表'
                    )
                    pd.DataFrame({
                        '卡方统计量': [results['chi_square']['chi2_statistic']],
                        'P值': [results['chi_square']['p_value']],
                        '自由度': [results['chi_square']['degrees_of_freedom']],
                        '克拉默系数': [results['chi_square']['cramers_v']]
                    }).to_excel(writer, sheet_name='卡方检验_结果')
                
                # 导出聚类分析结果
                if 'cluster' in results:
                    results['cluster']['cluster_centers'].to_excel(
                        writer,
                        sheet_name='聚类分析_中心点'
                    )
                    pd.DataFrame({
                        '轮廓系数': [results['cluster']['silhouette_score']]
                    }).to_excel(writer, sheet_name='聚类分析_评估')
                    results['cluster']['cluster_sizes'].to_excel(
                        writer,
                        sheet_name='聚类分析_样本分布'
                    )
                
        except Exception as e:
            logger.error(f"导出分析结果失败: {str(e)}")
            raise AnalysisError(f"无法导出分析结果: {str(e)}")

    def generate_report(self, results: Dict[str, Any], output_path: str) -> None:
        """生成分析报告（Word格式）"""
        try:
            from docx import Document
            doc = Document()
            
            # 添加标题
            doc.add_heading('统计分析报告', 0)
            
            # 添加描述性统计结果
            if 'descriptive' in results:
                doc.add_heading('描述性统计分析', level=1)
                # ...添加描述性统计结果的详细内容
            
            # 添加其他分析结果...
            
            # 保存文档
            doc.save(output_path)
            
        except Exception as e:
            logger.error(f"生成分析报告失败: {str(e)}")
            raise AnalysisError(f"无法生成分析报告: {str(e)}")