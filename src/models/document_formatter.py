from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from utils.exceptions import FormattingError
from utils.logger import setup_logger

logger = setup_logger(__name__)

class DocumentFormatter:
    def __init__(self, template_path: Optional[str] = None):
        self.doc = Document(template_path) if template_path else Document()
        self.styles = self._load_default_styles()
    
    def format_document(self, content: Dict[str, Any], format_type: str) -> Document:
        """格式化文档"""
        try:
            if format_type == 'thesis':
                return self._format_thesis(content)
            elif format_type == 'paper':
                return self._format_paper(content)
            else:
                raise FormattingError(f"不支持的文档类型: {format_type}")
        except Exception as e:
            logger.error(f"文档格式化失败: {str(e)}")
            raise FormattingError(f"无法格式化文档: {str(e)}")
    
    def _format_thesis(self, content: Dict[str, Any]) -> Document:
        """论文格式化"""
        try:
            # 设置页面格式
            section = self.doc.sections[0]
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            
            # 添加标题
            title = self.doc.add_heading(content['title'], level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加摘要
            self.doc.add_heading('摘要', level=1)
            abstract = self.doc.add_paragraph(content['abstract'])
            abstract.style = self.styles['abstract']
            
            # 添加正文
            for chapter in content['chapters']:
                self._add_chapter(chapter)
            
            return self.doc
            
        except Exception as e:
            logger.error(f"论文格式化失败: {str(e)}")
            raise FormattingError(f"无法格式化论文: {str(e)}")
    
    def _load_default_styles(self) -> Dict[str, Any]:
        """加载默认样式"""
        styles = {
            'title': {
                'name': '标题',
                'font_size': Pt(18),
                'bold': True,
                'alignment': WD_PARAGRAPH_ALIGNMENT.CENTER
            },
            'heading1': {
                'name': '一级标题',
                'font_size': Pt(16),
                'bold': True,
                'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT
            },
            'heading2': {
                'name': '二级标题',
                'font_size': Pt(14),
                'bold': True,
                'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT
            },
            'normal': {
                'name': '正文',
                'font_size': Pt(12),
                'line_spacing': 1.5,
                'alignment': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            },
            'abstract': {
                'name': '摘要',
                'font_size': Pt(12),
                'italic': True,
                'alignment': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
        }
        
        # 应用样式到文档
        for style_name, props in styles.items():
            try:
                style = self.doc.styles.add_style(props['name'], 1)
                font = style.font
                font.size = props['font_size']
                font.bold = props.get('bold', False)
                font.italic = props.get('italic', False)
                
                paragraph_format = style.paragraph_format
                paragraph_format.alignment = props['alignment']
                if 'line_spacing' in props:
                    paragraph_format.line_spacing = props['line_spacing']
                
            except Exception as e:
                logger.warning(f"样式 {style_name} 创建失败: {str(e)}")
                
        return styles

class ReferenceManager:
    def __init__(self, style: str = 'GB/T 7714'):
        self.style = style
        self.references = []
        
    def add_reference(self, reference: Dict[str, str]):
        """添加参考文献"""
        self.references.append(reference)
        
    def format_citation(self, ref_id: str) -> str:
        """格式化引用标记"""
        ref = self._find_reference(ref_id)
        if self.style == 'GB/T 7714':
            return f"[{self.references.index(ref) + 1}]"
        elif self.style == 'APA':
            return f"({ref['author']}, {ref['year']})"
        
    def generate_bibliography(self) -> List[str]:
        """生成参考文献列表"""
        formatted_refs = []
        
        for ref in self.references:
            if self.style == 'GB/T 7714':
                formatted = self._format_gbt7714(ref)
            elif self.style == 'APA':
                formatted = self._format_apa(ref)
            formatted_refs.append(formatted)
            
        return formatted_refs
    
    def _format_gbt7714(self, ref: Dict[str, str]) -> str:
        """GB/T 7714格式化"""
        if ref['type'] == 'journal':
            return (f"{ref['author']}. {ref['title']}[J]. "
                   f"{ref['journal']}, {ref['year']}, {ref['volume']}({ref['issue']}): "
                   f"{ref['pages']}")
        elif ref['type'] == 'book':
            return (f"{ref['author']}. {ref['title']}[M]. "
                   f"{ref['publisher']}, {ref['year']}")
        return ""