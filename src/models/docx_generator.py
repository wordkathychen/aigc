"""
Word文档生成模块
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging
from src.utils.logger import setup_logger

logger = setup_logger(__name__)

class DocxGenerator:
    """Word文档生成器"""
    
    def __init__(self, template_path=None):
        """初始化Word文档生成器
        
        Args:
            template_path: Word模板路径，如果为None则创建空白文档
        """
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
            logger.info(f"使用模板创建Word文档: {template_path}")
        else:
            self.doc = Document()
            logger.info("创建空白Word文档")
            
        # 设置默认字体
        self._set_default_font()
    
    def _set_default_font(self):
        """设置默认字体为宋体"""
        # 设置默认样式
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def generate_from_markdown(self, markdown_text, output_path):
        """从Markdown文本生成Word文档
        
        Args:
            markdown_text: Markdown格式的文本
            output_path: 输出文件路径
            
        Returns:
            str: 输出文件路径
        """
        try:
            # 解析Markdown文本
            sections = self._parse_markdown(markdown_text)
            
            # 生成文档内容
            for section in sections:
                self._add_section_to_doc(section)
            
            # 保存文档
            self.doc.save(output_path)
            logger.info(f"Word文档已保存: {output_path}")
            
            return output_path
        
        except Exception as e:
            logger.error(f"生成Word文档失败: {str(e)}")
            raise
    
    def _parse_markdown(self, markdown_text):
        """解析Markdown文本，提取章节结构
        
        Args:
            markdown_text: Markdown格式的文本
            
        Returns:
            List: 章节列表，每个章节包含标题级别、标题文本和内容
        """
        lines = markdown_text.split('\n')
        sections = []
        current_section = None
        
        for line in lines:
            # 检查是否是标题行
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            
            if header_match:
                # 如果已有当前章节，保存它
                if current_section:
                    sections.append(current_section)
                
                # 创建新章节
                level = len(header_match.group(1))
                title = header_match.group(2).strip()
                current_section = {
                    'level': level,
                    'title': title,
                    'content': []
                }
            elif current_section:
                # 添加内容到当前章节
                current_section['content'].append(line)
            else:
                # 文档开始的内容，创建一个无标题章节
                current_section = {
                    'level': 0,
                    'title': '',
                    'content': [line]
                }
        
        # 添加最后一个章节
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _add_section_to_doc(self, section):
        """将章节添加到Word文档
        
        Args:
            section: 章节信息，包含级别、标题和内容
        """
        level = section['level']
        title = section['title']
        content = '\n'.join(section['content'])
        
        # 添加标题
        if level > 0 and title:
            # Word中的标题级别
            word_level = min(level, 9)  # Word最多支持9级标题
            
            # 添加标题段落
            heading = self.doc.add_heading(level=word_level)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # 设置标题文本
            run = heading.add_run(title)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            # 设置字体大小
            if level == 1:
                run.font.size = Pt(18)
                run.bold = True
            elif level == 2:
                run.font.size = Pt(16)
                run.bold = True
            elif level == 3:
                run.font.size = Pt(14)
                run.bold = True
            else:
                run.font.size = Pt(12)
                run.bold = True
        
        # 处理内容
        if content.strip():
            self._process_content(content)
    
    def _process_content(self, content):
        """处理Markdown内容，转换为Word格式
        
        Args:
            content: Markdown格式的内容
        """
        # 分段处理
        paragraphs = re.split(r'\n\s*\n', content)
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
                
            # 处理列表
            if re.match(r'^\s*[*\-+]\s', para_text):
                self._add_list(para_text)
                continue
            
            # 处理表格
            if '|' in para_text and re.search(r'\|[\-:]+\|', para_text):
                self._add_table(para_text)
                continue
            
            # 处理普通段落
            para = self.doc.add_paragraph()
            
            # 处理段落中的格式
            self._add_formatted_text(para, para_text)
    
    def _add_list(self, list_text):
        """添加列表
        
        Args:
            list_text: 列表文本
        """
        lines = list_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 移除列表标记
            match = re.match(r'^\s*[*\-+]\s+(.+)$', line)
            if match:
                text = match.group(1)
                
                # 添加项目符号
                item = self.doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(item, text)
    
    def _add_table(self, table_text):
        """添加表格
        
        Args:
            table_text: Markdown表格文本
        """
        lines = table_text.split('\n')
        rows = []
        
        # 解析表格行
        for line in lines:
            line = line.strip()
            if not line or line.startswith('|---'):
                continue
                
            # 解析单元格
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            rows.append(cells)
        
        if not rows:
            return
            
        # 创建表格
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        
        # 填充表格内容
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    para = cell.paragraphs[0]
                    self._add_formatted_text(para, cell_text)
    
    def _add_formatted_text(self, paragraph, text):
        """添加带格式的文本
        
        Args:
            paragraph: 段落对象
            text: 文本内容
        """
        # 处理粗体
        bold_parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in bold_parts:
            if part.startswith('**') and part.endswith('**'):
                # 粗体文本
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # 处理斜体
                italic_parts = re.split(r'(\*.*?\*)', part)
                
                for italic_part in italic_parts:
                    if italic_part.startswith('*') and italic_part.endswith('*'):
                        # 斜体文本
                        run = paragraph.add_run(italic_part[1:-1])
                        run.italic = True
                    else:
                        # 普通文本
                        if italic_part:
                            run = paragraph.add_run(italic_part)
                            
                            # 设置中文字体
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def add_cover_page(self, title, author, date, institution=None):
        """添加封面页
        
        Args:
            title: 标题
            author: 作者
            date: 日期
            institution: 机构名称
        """
        # 添加标题
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.name = '黑体'
        title_run.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 添加空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # 添加作者
        if author:
            author_para = self.doc.add_paragraph()
            author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            author_run = author_para.add_run(f"作者：{author}")
            author_run.font.size = Pt(14)
            author_run.font.name = '宋体'
            author_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加机构
        if institution:
            inst_para = self.doc.add_paragraph()
            inst_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            inst_run = inst_para.add_run(institution)
            inst_run.font.size = Pt(14)
            inst_run.font.name = '宋体'
            inst_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加日期
        if date:
            date_para = self.doc.add_paragraph()
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            date_run = date_para.add_run(date)
            date_run.font.size = Pt(14)
            date_run.font.name = '宋体'
            date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加分页符
        self.doc.add_page_break()
    
    def add_toc(self):
        """添加目录"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('目录')
        run.font.size = Pt(16)
        run.font.name = '黑体'
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 添加目录域代码
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        
        # 创建目录域
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "右键点击更新目录"
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        r_element = run._element
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)
        r_element.append(fldChar4)
        
        # 添加分页符
        self.doc.add_page_break() 