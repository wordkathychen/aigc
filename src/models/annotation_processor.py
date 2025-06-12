"""
批注处理模块
用于分析PDF中的批注并修改对应文本
"""

import os
import re
import fitz  # PyMuPDF
import docx
from docx.shared import RGBColor
from typing import List, Dict, Tuple, Optional, Any
import logging
from src.utils.logger import setup_logger

logger = setup_logger(__name__)

class AnnotationItem:
    """批注项"""
    def __init__(self, page_num: int, content: str, rect: List[float], text_context: str = ""):
        self.page_num = page_num  # 页码
        self.content = content    # 批注内容
        self.rect = rect          # 批注位置 [x0, y0, x1, y1]
        self.text_context = text_context  # 相关文本上下文
        self.replaced_text = ""   # 替换后的文本
        self.processed = False    # 是否已处理
        
    def __str__(self):
        return f"页码: {self.page_num+1}, 批注: {self.content}, 上下文: {self.text_context[:30]}..."

class AnnotationProcessor:
    """批注处理器"""
    
    def __init__(self, api_manager=None):
        """初始化批注处理器
        
        Args:
            api_manager: API管理器实例，用于生成替换文本
        """
        self.api_manager = api_manager
        self.annotations = []  # 批注列表
        self.pdf_text = {}     # PDF文本内容 {页码: 文本内容}
        self.current_file = None  # 当前处理的文件路径
    
    def extract_annotations(self, pdf_path: str) -> List[AnnotationItem]:
        """提取PDF中的批注
        
        Args:
            pdf_path: PDF文件路径
            
        Returns:
            批注列表
        """
        if not os.path.exists(pdf_path):
            logger.error(f"文件不存在: {pdf_path}")
            raise FileNotFoundError(f"文件不存在: {pdf_path}")
            
        self.annotations = []
        self.pdf_text = {}
        self.current_file = pdf_path
        
        try:
            # 打开PDF文件
            doc = fitz.open(pdf_path)
            
            # 遍历每一页
            for page_num, page in enumerate(doc):
                # 提取页面文本
                self.pdf_text[page_num] = page.get_text()
                
                # 提取批注
                annots = page.annots()
                if annots:
                    for annot in annots:
                        if annot.type[0] == 8:  # 注释类型
                            content = annot.info.get("content", "").strip()
                            if content:
                                # 获取批注位置
                                rect = list(annot.rect)
                                
                                # 获取批注附近的文本上下文
                                text_context = self._get_text_context(page, rect)
                                
                                # 创建批注项
                                annotation = AnnotationItem(
                                    page_num=page_num,
                                    content=content,
                                    rect=rect,
                                    text_context=text_context
                                )
                                self.annotations.append(annotation)
            
            logger.info(f"从PDF中提取了 {len(self.annotations)} 条批注")
            return self.annotations
        
        except Exception as e:
            logger.error(f"提取批注失败: {str(e)}")
            raise
    
    def _get_text_context(self, page, rect: List[float], context_range: float = 50.0) -> str:
        """获取批注附近的文本上下文
        
        Args:
            page: PDF页面
            rect: 批注位置 [x0, y0, x1, y1]
            context_range: 上下文范围（像素）
            
        Returns:
            文本上下文
        """
        # 扩大区域以包含更多上下文
        expanded_rect = [
            rect[0] - context_range,
            rect[1] - context_range,
            rect[2] + context_range,
            rect[3] + context_range
        ]
        
        # 获取扩大区域内的文本
        text = page.get_text("text", clip=expanded_rect)
        return text.strip()
    
    def process_annotations(self) -> List[AnnotationItem]:
        """处理批注，生成替换文本
        
        Returns:
            处理后的批注列表
        """
        if not self.annotations:
            logger.warning("没有批注需要处理")
            return []
        
        for i, annotation in enumerate(self.annotations):
            try:
                # 根据批注内容生成替换文本
                replaced_text = self._generate_replacement(annotation.content, annotation.text_context)
                
                # 设置替换后的文本
                annotation.replaced_text = replaced_text
                annotation.processed = True
                
                logger.info(f"处理批注 {i+1}/{len(self.annotations)}: {annotation.content[:30]}...")
            except Exception as e:
                logger.error(f"处理批注失败: {str(e)}")
                annotation.processed = False
        
        return self.annotations
    
    def _generate_replacement(self, annotation_content: str, text_context: str) -> str:
        """根据批注内容生成替换文本
        
        Args:
            annotation_content: 批注内容
            text_context: 文本上下文
            
        Returns:
            替换后的文本
        """
        # 如果有API管理器，使用LLM生成替换文本
        if self.api_manager:
            try:
                # 构建提示词
                prompt = f"""
                请根据以下批注修改文本内容:
                
                文本上下文: "{text_context}"
                
                批注要求: "{annotation_content}"
                
                请直接提供修改后的完整文本段落，不要添加任何解释。
                """
                
                # 获取LLM回复
                # 这里根据API管理器的实际接口调用方法进行修改
                response = self.api_manager.generate_content_sync(prompt, max_tokens=1000)
                
                # 清理回复，只保留文本内容
                replacement = response.strip()
                return replacement
            except Exception as e:
                logger.error(f"生成替换文本失败: {str(e)}")
                # 失败时返回原文
                return text_context
        
        # 没有API管理器时，使用简单的替换策略
        # 这是一个简化版本，实际应用中可能需要更复杂的处理
        replacement = text_context
        
        # 常见的批注模式，例如"将XX替换为YY"
        replace_pattern = r"将\s*[\"']?(.*?)[\"']?\s*替换为\s*[\"']?(.*?)[\"']?"
        replace_match = re.search(replace_pattern, annotation_content)
        if replace_match:
            old_text = replace_match.group(1)
            new_text = replace_match.group(2)
            replacement = text_context.replace(old_text, new_text)
        
        # 其他常见模式可以在这里添加
        
        return replacement
    
    def generate_corrected_document(self, output_path: str) -> str:
        """生成修正后的Word文档
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            输出文件的路径
        """
        if not self.annotations or not self.pdf_text:
            raise ValueError("没有可用的批注或PDF文本")
        
        # 创建新的Word文档
        doc = docx.Document()
        
        # 添加文档标题
        filename = os.path.basename(self.current_file)
        doc.add_heading(f"批注修正文档 - {filename}", 0)
        
        # 添加说明段落
        doc.add_paragraph("本文档包含根据PDF批注自动生成的修正文本。绿色文本表示已修改的部分。")
        
        # 处理每个批注
        for i, annotation in enumerate(self.annotations):
            # 添加批注信息
            doc.add_heading(f"批注 {i+1}: 第 {annotation.page_num + 1} 页", level=1)
            
            # 添加批注内容
            p = doc.add_paragraph("批注内容: ")
            p.add_run(annotation.content)
            
            # 添加原始上下文
            doc.add_heading("原始文本:", level=2)
            doc.add_paragraph(annotation.text_context)
            
            # 添加修改后文本
            doc.add_heading("修改后文本:", level=2)
            p = doc.add_paragraph()
            
            # 设置修改后文本为绿色
            if annotation.processed and annotation.replaced_text:
                run = p.add_run(annotation.replaced_text)
                run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            else:
                p.add_run("(未能生成修改文本)")
            
            # 添加分隔线
            doc.add_paragraph("=" * 50)
        
        # 保存文档
        doc.save(output_path)
        logger.info(f"修正文档已保存到: {output_path}")
        
        return output_path
    
    def extract_text_with_annotations(self) -> Dict[int, List[Dict[str, Any]]]:
        """提取带有批注位置的文本
        
        Returns:
            按页码组织的文本和批注信息
            {页码: [{文本, 是否有批注, 批注内容, 替换文本}, ...]}
        """
        if not self.annotations or not self.pdf_text:
            return {}
        
        result = {}
        
        # 遍历每一页
        for page_num, page_text in self.pdf_text.items():
            result[page_num] = []
            
            # 获取当前页的批注
            page_annotations = [a for a in self.annotations if a.page_num == page_num]
            
            if not page_annotations:
                # 如果没有批注，直接添加整页文本
                result[page_num].append({
                    'text': page_text,
                    'has_annotation': False,
                    'annotation': None,
                    'replacement': None
                })
                continue
            
            # TODO: 这里需要更复杂的文本分段和批注匹配逻辑
            # 简化版：将每个批注视为独立段落
            for annotation in page_annotations:
                result[page_num].append({
                    'text': annotation.text_context,
                    'has_annotation': True,
                    'annotation': annotation.content,
                    'replacement': annotation.replaced_text if annotation.processed else None
                })
        
        return result

# 单例模式
annotation_processor = AnnotationProcessor() 