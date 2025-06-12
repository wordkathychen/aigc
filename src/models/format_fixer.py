"""
格式修复模块
用于处理文档格式的统一化和修复
"""

import os
import tempfile
import shutil
from typing import Dict, Any, Optional
import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

from src.utils.logger import setup_logger

logger = setup_logger(__name__)

class FormatFixer:
    """文档格式修复类"""
    
    def __init__(self):
        """初始化格式修复器"""
        self.temp_dir = tempfile.mkdtemp()
        logger.info(f"初始化格式修复器，临时目录: {self.temp_dir}")
        
    def __del__(self):
        """析构函数，清理临时文件"""
        try:
            shutil.rmtree(self.temp_dir)
            logger.info("清理格式修复器临时文件")
        except Exception as e:
            logger.error(f"清理临时文件失败: {str(e)}")
    
    def fix_format(self, source_file: str, template_file: str, options: Dict[str, bool]) -> Optional[str]:
        """修复文档格式
        
        Args:
            source_file: 源文件路径
            template_file: 模板文件路径
            options: 格式调整选项
                - adjust_fonts: 是否调整字体
                - adjust_spacing: 是否调整行距和段落间距
                - adjust_margins: 是否调整页边距
                - adjust_headings: 是否调整标题样式
                - keep_content: 是否保留原文内容
                
        Returns:
            str: 处理后的文件路径，失败返回None
        """
        if not os.path.exists(source_file):
            logger.error(f"源文件不存在: {source_file}")
            return None
            
        if not os.path.exists(template_file):
            logger.error(f"模板文件不存在: {template_file}")
            return None
            
        try:
            # 加载源文档和模板文档
            source_doc = docx.Document(source_file)
            template_doc = docx.Document(template_file)
            
            # 创建新文档
            new_doc = docx.Document()
            
            # 复制模板文档的页面设置
            self._copy_page_setup(template_doc, new_doc)
            
            # 复制模板文档的样式
            self._copy_styles(template_doc, new_doc, options)
            
            # 复制源文档的内容，应用模板样式
            self._copy_content(source_doc, new_doc, template_doc, options)
            
            # 保存结果
            result_path = os.path.join(self.temp_dir, "formatted_document.docx")
            new_doc.save(result_path)
            
            logger.info(f"格式修复完成，结果保存至: {result_path}")
            return result_path
            
        except Exception as e:
            logger.error(f"格式修复失败: {str(e)}")
            return None
    
    def _copy_page_setup(self, template_doc: docx.Document, new_doc: docx.Document) -> None:
        """复制页面设置
        
        Args:
            template_doc: 模板文档
            new_doc: 新文档
        """
        # 复制页面设置
        for section in template_doc.sections:
            new_section = new_doc.sections[0]  # 文档至少有一个section
            
            # 页面大小
            new_section.page_height = section.page_height
            new_section.page_width = section.page_width
            
            # 页边距
            new_section.left_margin = section.left_margin
            new_section.right_margin = section.right_margin
            new_section.top_margin = section.top_margin
            new_section.bottom_margin = section.bottom_margin
            
            # 页眉页脚
            new_section.header_distance = section.header_distance
            new_section.footer_distance = section.footer_distance
            
            # 纸张方向
            new_section.orientation = section.orientation
            
            # 一般只复制第一个section的设置
            break
    
    def _copy_styles(self, template_doc: docx.Document, new_doc: docx.Document, options: Dict[str, bool]) -> None:
        """复制样式
        
        Args:
            template_doc: 模板文档
            new_doc: 新文档
            options: 格式调整选项
        """
        # 复制样式
        if options.get('adjust_fonts', True):
            # 复制默认字体
            new_doc.styles['Normal'].font.name = template_doc.styles['Normal'].font.name
            new_doc.styles['Normal'].font.size = template_doc.styles['Normal'].font.size
            
        if options.get('adjust_spacing', True):
            # 复制默认段落间距
            if hasattr(template_doc.styles['Normal'].paragraph_format, 'space_before'):
                new_doc.styles['Normal'].paragraph_format.space_before = template_doc.styles['Normal'].paragraph_format.space_before
            
            if hasattr(template_doc.styles['Normal'].paragraph_format, 'space_after'):
                new_doc.styles['Normal'].paragraph_format.space_after = template_doc.styles['Normal'].paragraph_format.space_after
            
            # 复制默认行距
            if hasattr(template_doc.styles['Normal'].paragraph_format, 'line_spacing'):
                new_doc.styles['Normal'].paragraph_format.line_spacing = template_doc.styles['Normal'].paragraph_format.line_spacing
            
        if options.get('adjust_headings', True):
            # 复制标题样式
            heading_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5']
            for style_name in heading_styles:
                if style_name in template_doc.styles and style_name in new_doc.styles:
                    # 复制字体
                    new_doc.styles[style_name].font.name = template_doc.styles[style_name].font.name
                    new_doc.styles[style_name].font.size = template_doc.styles[style_name].font.size
                    new_doc.styles[style_name].font.bold = template_doc.styles[style_name].font.bold
                    
                    # 复制段落格式
                    if hasattr(template_doc.styles[style_name].paragraph_format, 'space_before'):
                        new_doc.styles[style_name].paragraph_format.space_before = template_doc.styles[style_name].paragraph_format.space_before
                    
                    if hasattr(template_doc.styles[style_name].paragraph_format, 'space_after'):
                        new_doc.styles[style_name].paragraph_format.space_after = template_doc.styles[style_name].paragraph_format.space_after
    
    def _copy_content(self, source_doc: docx.Document, new_doc: docx.Document, template_doc: docx.Document, options: Dict[str, bool]) -> None:
        """复制内容
        
        Args:
            source_doc: 源文档
            new_doc: 新文档
            template_doc: 模板文档
            options: 格式调整选项
        """
        if options.get('keep_content', True):
            # 复制源文档的内容，应用模板样式
            for paragraph in source_doc.paragraphs:
                # 创建新段落
                new_para = new_doc.add_paragraph()
                
                # 复制文本
                new_para.text = paragraph.text
                
                # 应用样式
                if paragraph.style.name.startswith('Heading'):
                    new_para.style = paragraph.style.name
                else:
                    new_para.style = 'Normal'
                    
                # 应用格式调整
                if options.get('adjust_fonts', True):
                    for run in new_para.runs:
                        run.font.name = template_doc.styles['Normal'].font.name
                        
                if options.get('adjust_spacing', True):
                    if hasattr(template_doc.styles['Normal'].paragraph_format, 'space_before'):
                        new_para.paragraph_format.space_before = template_doc.styles['Normal'].paragraph_format.space_before
                    
                    if hasattr(template_doc.styles['Normal'].paragraph_format, 'space_after'):
                        new_para.paragraph_format.space_after = template_doc.styles['Normal'].paragraph_format.space_after
                    
                    if hasattr(template_doc.styles['Normal'].paragraph_format, 'line_spacing'):
                        new_para.paragraph_format.line_spacing = template_doc.styles['Normal'].paragraph_format.line_spacing
            
            # 复制表格
            for table in source_doc.tables:
                new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                
                # 复制表格内容
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text
                
                # 应用表格样式
                new_table.style = 'Table Grid'
        else:
            # 使用模板文档的内容结构
            for paragraph in template_doc.paragraphs:
                # 创建新段落
                new_para = new_doc.add_paragraph()
                
                # 复制文本
                new_para.text = paragraph.text
                
                # 应用样式
                new_para.style = paragraph.style.name
            
            # 复制表格
            for table in template_doc.tables:
                new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                
                # 复制表格内容
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text
                
                # 应用表格样式
                new_table.style = table.style 