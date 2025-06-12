import os
import io
import tempfile
import re
from typing import Dict, List, Optional, Union, Tuple
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import PyPDF2
import pandas as pd
from utils.exceptions import FileHandlerError
from utils.logger import setup_logger
from config.settings import ALLOWED_FILE_TYPES

logger = setup_logger(__name__)

class FileHandler:
    """文件读取和保存处理类"""
    
    def __init__(self):
        """初始化文件处理器"""
        # 支持的文件类型
        self.supported_extensions = {
            'text': ['.txt', '.md'],
            'word': ['.docx', '.doc'],
            'pdf': ['.pdf'],
            'excel': ['.xlsx', '.xls', '.csv'],
            'ppt': ['.pptx', '.ppt']
        }
    
    def read_file(self, filepath: str) -> str:
        """读取文件内容
        
        参数:
            filepath: 文件路径
            
        返回:
            str: 文件内容
        """
        _, ext = os.path.splitext(filepath)
        ext = ext.lower()
        
        try:
            # 基于文件扩展名选择适当的读取方法
            if ext in self.supported_extensions['text']:
                return self._read_text_file(filepath)
            elif ext in self.supported_extensions['word']:
                return self._read_word_file(filepath)
            elif ext in self.supported_extensions['pdf']:
                return self._read_pdf_file(filepath)
            elif ext in self.supported_extensions['excel']:
                return self._read_excel_file(filepath)
            else:
                raise ValueError(f"不支持的文件类型: {ext}")
        except Exception as e:
            logger.error(f"读取文件失败: {str(e)}")
            raise
    
    def save_file(self, filepath: str, content: str, 
                 original_content: Optional[str] = None,
                 highlight_changes: bool = False) -> bool:
        """保存内容到文件
        
        参数:
            filepath: 文件保存路径
            content: 要保存的内容
            original_content: 原始内容（用于比较和高亮变化）
            highlight_changes: 是否高亮显示变化
            
        返回:
            bool: 是否保存成功
        """
        try:
            _, ext = os.path.splitext(filepath)
            ext = ext.lower()
            
            # 基于文件扩展名选择适当的保存方法
            if ext in self.supported_extensions['text']:
                return self._save_text_file(filepath, content)
            elif ext in self.supported_extensions['word']:
                return self._save_word_file(filepath, content, 
                                         original_content, highlight_changes)
            elif ext in self.supported_extensions['excel']:
                return self._save_excel_file(filepath, content)
            else:
                raise ValueError(f"不支持的文件类型: {ext}")
        except Exception as e:
            logger.error(f"保存文件失败: {str(e)}")
            return False
    
    def _read_text_file(self, filepath: str) -> str:
        """读取文本文件"""
        if not os.path.exists(filepath):
            raise FileHandlerError(f"文件不存在: {filepath}")
            
        try:
            # 自动检测文件编码
            with open(filepath, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
            
            with open(filepath, 'r', encoding=encoding) as f:
                return f.read()
        except Exception as e:
            logger.error(f"读取文本文件失败: {str(e)}")
            raise FileHandlerError(f"无法读取文件 {filepath}: {str(e)}")

    def _read_word_file(self, filepath: str) -> str:
        """读取Word文档"""
        if not filepath.endswith(('.docx', '.doc')):
            raise FileHandlerError("不是有效的Word文档格式")
            
        try:
            doc = docx.Document(filepath)
            content = []
            
            # 提取正文
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text)
            
            return "\n".join(content)
        except Exception as e:
            logger.error(f"读取Word文档失败: {str(e)}")
            raise FileHandlerError(f"无法读取文件 {filepath}: {str(e)}")

    def _read_pdf_file(self, filepath: str) -> str:
        """读取PDF文件"""
        if not filepath.endswith('.pdf'):
            raise FileHandlerError("不是有效的PDF文件格式")
            
        try:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content.append(text)
                
                return "\n".join(content)
        except Exception as e:
            logger.error(f"读取PDF文件失败: {str(e)}")
            raise FileHandlerError(f"无法读取文件 {filepath}: {str(e)}")

    def _read_excel_file(self, filepath: str) -> str:
        """读取Excel文件并转换为文本"""
        _, ext = os.path.splitext(filepath)
        if ext.lower() == '.csv':
            df = pd.read_csv(filepath)
        else:
            df = pd.read_excel(filepath)
        
        # 将DataFrame转换为字符串
        buffer = io.StringIO()
        df.to_csv(buffer, index=False)
        return buffer.getvalue()
    
    def _save_text_file(self, filepath: str, content: str) -> bool:
        """保存文本文件"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return True
    
    def _save_word_file(self, filepath: str, content: str, 
                      original_content: Optional[str] = None,
                      highlight_changes: bool = False) -> bool:
        """保存Word文档"""
        try:
            doc = docx.Document()
            
            if highlight_changes and original_content:
                # 比较文本并高亮差异
                paragraphs = content.split('\n')
                original_paragraphs = original_content.split('\n')
                
                # 创建一个映射，跟踪哪些段落有变化
                changes = {}
                
                for i, para in enumerate(paragraphs):
                    if i < len(original_paragraphs):
                        original_para = original_paragraphs[i]
                        # 简单比较，实际应用可能需要更复杂的diff算法
                        if para != original_para:
                            changes[i] = self._find_changes(original_para, para)
                    else:
                        # 新增段落
                        changes[i] = [(0, len(para))]
                
                # 添加带高亮的段落
                for i, para_text in enumerate(paragraphs):
                    p = doc.add_paragraph()
                    if i in changes:
                        # 段落有变化，需要高亮
                        current_pos = 0
                        for start, end in changes[i]:
                            # 添加未变化部分
                            if start > current_pos:
                                p.add_run(para_text[current_pos:start])
                            
                            # 添加高亮部分
                            highlighted_text = para_text[start:end]
                            run = p.add_run(highlighted_text)
                            run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                            
                            current_pos = end
                        
                        # 添加剩余部分
                        if current_pos < len(para_text):
                            p.add_run(para_text[current_pos:])
                    else:
                        # 段落无变化
                        p.add_run(para_text)
            else:
                # 不需要高亮，直接添加段落
                for para in content.split('\n'):
                    if para.strip():  # 跳过空段落
                        doc.add_paragraph(para)
            
            # 确保目录存在
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # 保存文件
            doc.save(filepath)
            return True
        except Exception as e:
            logger.error(f"保存Word文档失败: {str(e)}")
            raise FileHandlerError(f"无法保存文件 {filepath}: {str(e)}")
    
    def _save_excel_file(self, filepath: str, content: str) -> bool:
        """保存Excel文件"""
        try:
            # 尝试解析内容为DataFrame
            df = pd.read_csv(io.StringIO(content))
            
            # 基于文件扩展名选择保存格式
            _, ext = os.path.splitext(filepath)
            if ext.lower() == '.csv':
                df.to_csv(filepath, index=False)
            else:
                df.to_excel(filepath, index=False)
            return True
        except Exception as e:
            logger.error(f"保存Excel文件失败: {str(e)}")
            return False
    
    def _find_changes(self, original: str, modified: str) -> List[Tuple[int, int]]:
        """查找文本中的变化部分
        
        返回一个列表，包含(开始位置, 结束位置)元组
        """
        import difflib
        
        # 使用difflib查找差异
        matcher = difflib.SequenceMatcher(None, original, modified)
        changes = []
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag in ('replace', 'insert'):
                changes.append((j1, j2))
        
        return changes
    
    def get_file_type(self, filepath: str) -> str:
        """获取文件类型"""
        _, ext = os.path.splitext(filepath)
        ext = ext.lower()
        
        for file_type, extensions in self.supported_extensions.items():
            if ext in extensions:
                return file_type
        
        return "unknown"

    @staticmethod
    def validate_file(file_path: str, allowed_types: Optional[List[str]] = None) -> None:
        """验证文件是否有效"""
        if not os.path.exists(file_path):
            raise FileHandlerError("文件不存在")
            
        if not os.path.isfile(file_path):
            raise FileHandlerError("不是有效的文件")
            
        if allowed_types:
            ext = os.path.splitext(file_path)[1].lower()
            if ext not in allowed_types:
                raise FileHandlerError(f"不支持的文件类型: {ext}")